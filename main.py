import sys #Built-in for Python 3.12.6
import os #Built-in for Python 3.12.6
import ocrmypdf #v16.5.0
import fitz  # Also know as PyMuPDF, v1.24.10 
from docx import Document
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QPushButton, QLabel,
    QVBoxLayout, QHBoxLayout, QWidget, QMessageBox, QComboBox, QTextEdit, QScrollArea
) #5.15.11 
from PyQt5.QtGui import QImage, QPixmap #v5.15.11 
from PyQt5.QtCore import Qt #v5.15.11 


class OCRApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.init_ui()
        self.selected_files = []  # To store selected files
        self.output_path = None
        self.mode = 'single'  # 'single' or 'multiple'

    def init_ui(self):
        self.setWindowTitle('OCR My PDF')
        self.setGeometry(100, 100, 1600, 800)

        # Create widgets
        self.label_input = QLabel('Input Path: Not Selected', self)
        self.label_output = QLabel('Output Path: Not Selected', self)

        # Dropdown for choosing single or multiple files
        self.combo_mode = QComboBox(self)
        self.combo_mode.addItems(['Single File', 'Multiple Files'])
        self.combo_mode.currentIndexChanged.connect(self.change_mode)

        self.btn_upload = QPushButton('Upload PDF', self)
        self.btn_upload.clicked.connect(self.upload_files)

        self.btn_output = QPushButton('Select Output Path', self)
        self.btn_output.clicked.connect(self.select_output_path)

        self.btn_process = QPushButton('Process PDF', self)
        self.btn_process.clicked.connect(self.process_files)

        # TextEdit fields for comparison (OCR PDF and DOCX)
        self.ocr_text_edit = QTextEdit(self)
        self.docx_text_edit = QTextEdit(self)
        self.ocr_text_edit.setReadOnly(True)
        self.docx_text_edit.setReadOnly(True)

        # Create a scroll area for image-based PDFs
        self.scroll_area = QScrollArea(self)
        self.scroll_area.setWidgetResizable(True)

        # Set up the layout for displaying images
        self.scroll_area_widget_layout = QVBoxLayout()
        self.scroll_area_widget = QWidget()
        self.scroll_area_widget.setLayout(self.scroll_area_widget_layout)
        self.scroll_area.setWidget(self.scroll_area_widget)

        # Layout
        main_layout = QVBoxLayout()
        main_layout.addWidget(self.label_input)
        main_layout.addWidget(self.combo_mode)
        main_layout.addWidget(self.btn_upload)
        main_layout.addWidget(self.label_output)
        main_layout.addWidget(self.btn_output)
        main_layout.addWidget(self.btn_process)

        # Horizontal layout for image display and text comparison
        horizontal_layout = QHBoxLayout()
        horizontal_layout.addWidget(self.scroll_area)

        # Add text comparison widgets side by side
        comparison_layout = QHBoxLayout()
        comparison_layout.addWidget(QLabel('Original PDF Text', self))
        comparison_layout.addWidget(self.ocr_text_edit)
        comparison_layout.addWidget(QLabel('DOCX Text', self))
        comparison_layout.addWidget(self.docx_text_edit)

        horizontal_layout.addLayout(comparison_layout)

        main_layout.addLayout(horizontal_layout)

        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)

    def change_mode(self):
        self.mode = 'multiple' if self.combo_mode.currentIndex() == 1 else 'single'
        self.selected_files = []  # Clear selected files when changing mode
        self.label_input.setText('Input Path: Not Selected')

    def upload_files(self):
        if self.mode == 'single':
            options = QFileDialog.Options()
            file, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf);;All Files (*)", options=options)
            if file:
                self.selected_files = [file]
                self.label_input.setText(f'Input Path: {file}')
                self.btn_output.setEnabled(True)
                self.btn_process.setEnabled(True)
        else:
            options = QFileDialog.Options()
            folder = QFileDialog.getExistingDirectory(self, "Select Folder", options=options)
            if folder:
                self.selected_files = [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith('.pdf')]
                if self.selected_files:
                    self.label_input.setText(f'Input Path: {len(self.selected_files)} files selected')
                    self.btn_output.setEnabled(True)
                    self.btn_process.setEnabled(True)

    def select_output_path(self):
        options = QFileDialog.Options()
        self.output_path = QFileDialog.getExistingDirectory(self, "Select Output Directory", options=options)
        if self.output_path:
            self.label_output.setText(f'Output Path: {self.output_path}')

    def process_files(self):
        if self.selected_files and self.output_path:
            for selected_file in self.selected_files:
                base_name = os.path.splitext(os.path.basename(selected_file))[0]
                output_dir = os.path.join(self.output_path, base_name)
                ocr_pdf_dir = os.path.join(output_dir, 'ocr_pdf')
                text_dir = os.path.join(output_dir, 'text')
                os.makedirs(ocr_pdf_dir, exist_ok=True)
                os.makedirs(text_dir, exist_ok=True)
                output_file = os.path.join(ocr_pdf_dir, os.path.basename(selected_file))

                try:
                    if self.is_text_based(selected_file):
                        # Directly split into chapters and convert to DOCX
                        self.split_into_chapters(selected_file, os.path.join(output_dir, 'chapters'))
                        self.convert_to_docx(selected_file, os.path.join(text_dir, f'{base_name}.docx'))
                    else:
                        # Perform OCR, split into chapters, and convert OCR PDF to DOCX
                        ocrmypdf.ocr(selected_file, output_file)
                        self.split_into_chapters(output_file, os.path.join(output_dir, 'chapters'))
                        self.convert_to_docx(output_file, os.path.join(text_dir, f'{base_name}.docx'))

                    # After processing, show the visual comparison
                    self.compare_and_show(selected_file, output_file, os.path.join(text_dir, f'{base_name}.docx'))

                except Exception as e:
                    print(f'Error processing file {selected_file}: {e}')

            self.show_message('Success', 'All files have been processed successfully.')

    def is_text_based(self, pdf_path):
        """ Check if the PDF contains text or is image-based """
        try:
            pdf_document = fitz.open(pdf_path)
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text("text")
                if text.strip():  # If there's any text on the page
                    pdf_document.close()
                    return True
            pdf_document.close()
            return False
        except Exception as e:
            print(f'Error checking PDF type: {e}')
            return False

    def split_into_chapters(self, pdf_path, chapters_dir):
        try:
            pdf_document = fitz.open(pdf_path)
            chapter_start_pages = []

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text("text")

                if text:
                    first_word = text.split()[0].lower()
                    if first_word == 'chapter':
                        chapter_start_pages.append(page_num)

            os.makedirs(chapters_dir, exist_ok=True)

            for i, start_page in enumerate(chapter_start_pages):
                end_page = chapter_start_pages[i + 1] if i + 1 < len(chapter_start_pages) else len(pdf_document)
                chapter_pdf_path = os.path.join(chapters_dir, f'{os.path.basename(pdf_path).replace(".pdf", f"_chapter_{i+1}.pdf")}')
                chapter_pdf = fitz.open()
                for page_num in range(start_page, end_page):
                    chapter_pdf.insert_pdf(pdf_document, from_page=page_num, to_page=page_num)
                chapter_pdf.save(chapter_pdf_path)
                chapter_pdf.close()

            pdf_document.close()
        except Exception as e:
            print(f'Error splitting chapters: {e}')

    def convert_to_docx(self, pdf_path, docx_output_path):
        """ Convert the PDF text to a DOCX file """
        try:
            pdf_document = fitz.open(pdf_path)
            doc = Document()

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text("text")
                doc.add_paragraph(text)

            doc.save(docx_output_path)
            pdf_document.close()
        except Exception as e:
            print(f'Error converting PDF to DOCX: {e}')

    def compare_and_show(self, original_pdf, ocr_pdf, docx_file):
        """Perform the visual comparison between the original, OCR, and DOCX files"""
        # If the original PDF is image-based, show the page images
        if not self.is_text_based(original_pdf):
            self.display_pdf_images(original_pdf)
        else:
            original_text = self.extract_text_from_pdf(original_pdf)
            self.ocr_text_edit.setText(original_text)

        ocr_text = self.extract_text_from_pdf(ocr_pdf)
        docx_text = self.extract_text_from_docx(docx_file)
        self.show_comparison(ocr_text, docx_text)

    def display_pdf_images(self, pdf_path):
        """ Display all page images of the original PDF in a scrollable view """
        self.scroll_area_widget_layout.setSpacing(10)
        self.scroll_area_widget_layout.setAlignment(Qt.AlignTop)

        pdf_document = fitz.open(pdf_path)
        self.scroll_area_widget_layout.addStretch()  # Add stretchable space before images

        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            pix = page.get_pixmap()

            # Convert PyMuPDF pixmap to QImage
            image = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format_RGB888)
            
            # Convert QImage to QPixmap
            qt_image = QPixmap.fromImage(image)
            
            # Create QLabel to hold the pixmap
            image_label = QLabel()
            image_label.setPixmap(qt_image)
            
            # Add the QLabel (image) to the layout
            self.scroll_area_widget_layout.addWidget(image_label)

        pdf_document.close()

    def extract_text_from_pdf(self, pdf_path):
        """ Extract text from a PDF file """
        try:
            pdf_document = fitz.open(pdf_path)
            text = ""
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text += page.get_text("text") + "\n"
            pdf_document.close()
            return text
        except Exception as e:
            print(f'Error extracting text from PDF: {e}')
            return ""

    def extract_text_from_docx(self, docx_path):
        """ Extract text from a DOCX file """
        try:
            doc = Document(docx_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f'Error extracting text from DOCX: {e}')
            return ""

    def show_comparison(self, text1, text2):
        """ Display comparison between two texts (OCR and DOCX) """
        self.ocr_text_edit.setText(text1)
        self.docx_text_edit.setText(text2)

    def show_message(self, title, message):
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle(title)
        msg_box.setText(message)
        msg_box.exec_()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = OCRApp()
    window.show()
    sys.exit(app.exec_())
